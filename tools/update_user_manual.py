"""Append the Phase C guide to the MolForge user manual."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path(__file__).resolve().parents[1]
MANUAL = ROOT / "MolForge_User_Manual.docx"
MARKER = "Phase C: Collaborative Platform"


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def main() -> None:
    document = Document(MANUAL)
    existing = "\n".join(paragraph.text for paragraph in document.paragraphs)
    if MARKER in existing:
        print("Phase C section already present.")
        return

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(MARKER, level=1)
    document.add_paragraph(
        "MolForge Phase C adds optional accounts, cloud-synced molecule libraries, public sharing, "
        "research workspaces, prediction feedback, and optional cloud calculations. Local editing "
        "and prediction continue to work without signing in."
    )

    document.add_heading("Accounts and Profiles", level=2)
    add_bullets(
        document,
        [
            "Select Sign in in the top navigation to use email/password, magic link, Google, or GitHub authentication.",
            "Open Profile from the user menu to update your research name and username and review daily usage.",
            "Free accounts receive 10 authenticated predictions per day. Higher tiers enable expanded compute and API access.",
        ],
    )

    document.add_heading("Save, Share, and Fork", level=2)
    add_bullets(
        document,
        [
            "In the Molecule Editor, select Save to store the molecule locally and sync it to Supabase when signed in.",
            "After a cloud save, select Share to make the molecule public, copy its link, or display its QR code.",
            "Anyone can open a public share link. Signed-in users can fork that molecule into their own library.",
            "Use Explore to search public molecules and sort by newest, most viewed, or most forked.",
        ],
    )

    document.add_heading("Workspaces and Comments", level=2)
    add_bullets(
        document,
        [
            "Open Workspaces to create a shared research space and invite collaborators by email.",
            "Workspace members share molecule references and receive Supabase Realtime updates.",
            "Public molecule pages include a discussion thread for signed-in collaborators.",
        ],
    )

    document.add_heading("Feedback and Native AI Learning", level=2)
    add_bullets(
        document,
        [
            "After a prediction, use the five-star accuracy control below the Properties panel.",
            "Expand Correct a value to submit a known experimental, literature, or external-tool result.",
            "Materials Project reconciliation and completed cloud calculations are stored privately as training feedback.",
            "Only administrators can read or export the accumulated feedback dataset.",
        ],
    )

    document.add_heading("Cloud Accuracy", level=2)
    add_bullets(
        document,
        [
            "Select Accurate in the editor to submit an optional xTB calculation for the active molecule.",
            "Open Cloud from the sidebar to inspect provider availability, cache statistics, and submit a SMILES calculation.",
            "Local xTB, Oracle, and GCP remain optional. MolForge queues requests gracefully when no provider is configured.",
            "The Native AI Colab notebook is stored at ml/training/MolForge_Native_AI_Training.ipynb.",
        ],
    )

    document.add_heading("Privacy and Configuration", level=2)
    document.add_paragraph(
        "MolForge uses Supabase Row Level Security so users can modify only their own private records. "
        "Public molecules are readable by anyone. Prediction feedback is private to administrators. "
        "Never place the Supabase service-role key in frontend/.env or commit any .env file."
    )

    document.save(MANUAL)
    print(f"Updated {MANUAL}")


if __name__ == "__main__":
    main()
